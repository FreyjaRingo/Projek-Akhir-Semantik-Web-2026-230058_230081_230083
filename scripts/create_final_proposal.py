from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "Proposal_Akhir_AnimeGraph_Nexus.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 32, 32)
MUTED = RGBColor(92, 92, 92)
LIGHT_FILL = "F4F6F9"
GRID = "D7DDE6"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("AnimeGraph Nexus | Proposal Akhir")
    set_run_font(run, size=9, color=MUTED)


def add_para(doc, text="", style=None, align=None, bold=False, italic=False, size=None, color=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        p.add_run(item)


def add_label_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [2700, 6660])
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
        set_cell_shading(table.cell(idx, 0), LIGHT_FILL)
        for cell in table.row_cells(idx):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=10.5, color=INK)
            table.cell(idx, 0).paragraphs[0].runs[0].bold = True
    add_para(doc, "", after=2)
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        set_cell_shading(cell, LIGHT_FILL)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, size=9.5, color=INK)
    set_table_width(table, widths)
    add_para(doc, "", after=2)
    return table


def cover(doc):
    add_para(doc, "PROPOSAL AKHIR", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=MUTED, size=12, after=10)
    add_para(doc, "AnimeGraph Nexus", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=26, color=INK, after=4)
    add_para(
        doc,
        "Platform eksplorasi anime berbasis Semantic Web, RDF, dan SPARQL",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        size=13,
        color=MUTED,
        after=24,
    )
    add_label_table(
        doc,
        [
            ("Mata Kuliah", "Semantic Web"),
            ("Nama Mahasiswa", "[Nama Mahasiswa]"),
            ("NIM", "[NIM]"),
            ("Program Studi", "[Program Studi]"),
            ("Dosen Pengampu", "[Nama Dosen Pengampu]"),
            ("Institusi", "[Nama Institusi]"),
            ("Tahun", "2026"),
        ],
    )
    add_para(
        doc,
        "Proposal ini disusun untuk menjelaskan rancangan, metodologi, implementasi, dan rencana evaluasi aplikasi AnimeGraph Nexus sebagai proyek akhir berbasis Semantic Web.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=MUTED,
        size=10.5,
        after=0,
    )
    doc.add_page_break()


def build_document():
    doc = Document()
    style_document(doc)
    cover(doc)

    add_heading(doc, "Ringkasan Proposal", 1)
    add_para(
        doc,
        "AnimeGraph Nexus adalah aplikasi eksplorasi data anime yang memanfaatkan pendekatan Semantic Web untuk merepresentasikan judul anime, genre, studio, tema, karakter, dan relasi antar karya dalam bentuk knowledge graph. Sistem ini mengolah data RDF lokal dan dataset eksternal menjadi struktur JSON yang dapat divisualisasikan melalui antarmuka React. Selain eksplorasi data lokal, sistem juga menyediakan route API server-side untuk melakukan pencarian data anime melalui Wikidata SPARQL endpoint.",
    )
    add_para(
        doc,
        "Fokus utama proposal ini adalah merancang aplikasi yang tidak hanya menampilkan daftar anime, tetapi juga memperlihatkan hubungan semantik antar entitas. Dengan demikian, pengguna dapat memahami keterkaitan antar judul anime berdasarkan jalur relasi seperti genre, studio produksi, tema naratif, karakter, dan rekomendasi terkait.",
    )

    add_heading(doc, "1. Latar Belakang", 1)
    add_para(
        doc,
        "Data anime umumnya tersedia dalam bentuk daftar atau katalog yang menampilkan atribut dasar seperti judul, genre, tahun rilis, dan studio. Bentuk penyajian tersebut berguna untuk pencarian sederhana, tetapi kurang mampu menunjukkan hubungan yang lebih luas antar entitas. Misalnya, dua anime dapat memiliki studio yang sama, genre yang serupa, tema naratif yang berdekatan, atau karakter dan karya terkait yang membentuk pola rekomendasi tertentu.",
    )
    add_para(
        doc,
        "Semantic Web menawarkan pendekatan yang lebih sesuai untuk persoalan tersebut karena data dapat dimodelkan sebagai resource, property, dan relation. Melalui RDF, informasi tidak hanya disimpan sebagai atribut terpisah, tetapi juga sebagai triple yang dapat menghubungkan satu entitas dengan entitas lain. SPARQL kemudian dapat digunakan untuk menelusuri hubungan tersebut secara deklaratif.",
    )
    add_para(
        doc,
        "Berdasarkan kebutuhan tersebut, AnimeGraph Nexus dirancang sebagai aplikasi web yang memadukan data RDF lokal, pemrosesan graph, pencarian semantik, visualisasi relasi, dan integrasi SPARQL. Aplikasi ini diharapkan dapat menunjukkan manfaat nyata Semantic Web dalam konteks eksplorasi data anime yang lebih interaktif dan mudah dipahami.",
    )

    add_heading(doc, "2. Rumusan Masalah", 1)
    add_numbered(
        doc,
        [
            "Bagaimana merepresentasikan data anime, genre, studio, tema, karakter, dan karya terkait sebagai knowledge graph berbasis RDF?",
            "Bagaimana mengolah data RDF dan dataset eksternal menjadi struktur yang siap digunakan oleh antarmuka web?",
            "Bagaimana menyediakan fitur pencarian, detail resource, perbandingan semantik, dan rekomendasi berbasis relasi?",
            "Bagaimana mengintegrasikan pencarian SPARQL melalui server-side API agar lebih aman dan stabil saat dideploy?",
            "Bagaimana menyajikan graph anime secara visual sehingga pengguna dapat memahami relasi antar entitas dengan mudah?",
        ],
    )

    add_heading(doc, "3. Tujuan", 1)
    add_bullets(
        doc,
        [
            "Membangun aplikasi web AnimeGraph Nexus sebagai media eksplorasi anime berbasis Semantic Web.",
            "Mengubah data RDF lokal dan dataset pendukung menjadi graph data yang dapat dicari, dibandingkan, dan divisualisasikan.",
            "Menyediakan fitur semantic search, resource detail, graph neighborhood, semantic compare, grounded QA, dan SPARQL query demo.",
            "Menyediakan serverless API route untuk pencarian data anime dari Wikidata melalui SPARQL.",
            "Menyusun antarmuka yang natural, rapi, dan sesuai dengan karakter aplikasi data exploration.",
        ],
    )

    add_heading(doc, "4. Manfaat", 1)
    add_heading(doc, "4.1 Manfaat Akademis", 2)
    add_bullets(
        doc,
        [
            "Memberikan contoh penerapan konsep RDF, relasi semantik, dan SPARQL dalam studi kasus yang konkret.",
            "Membantu memahami perbedaan antara penyajian data katalog biasa dan penyajian data berbasis knowledge graph.",
            "Menjadi bahan demonstrasi untuk menjelaskan query, resource, predicate, object, dan graph traversal pada mata kuliah Semantic Web.",
        ],
    )
    add_heading(doc, "4.2 Manfaat Praktis", 2)
    add_bullets(
        doc,
        [
            "Memudahkan pengguna mengeksplorasi hubungan antar anime berdasarkan studio, genre, tema, karakter, dan karya terkait.",
            "Menyediakan antarmuka yang lebih informatif dibandingkan daftar anime biasa karena relasi ditampilkan dalam bentuk graph.",
            "Menyediakan dasar pengembangan lanjutan untuk sistem rekomendasi anime berbasis relasi semantik.",
        ],
    )

    add_heading(doc, "5. Batasan Masalah", 1)
    add_bullets(
        doc,
        [
            "Data utama berasal dari RDF lokal dan dataset anime eksternal yang tersedia di lingkungan proyek.",
            "Relasi utama yang digunakan meliputi producedBy, hasGenre, hasTheme, featuresCharacter, dan relatedTo.",
            "Pencarian SPARQL eksternal difokuskan pada data anime dari Wikidata dengan batas hasil maksimal 20 item per query.",
            "Aplikasi dikembangkan sebagai web app berbasis React dan Vite, bukan sebagai sistem backend database penuh.",
            "Evaluasi sistem difokuskan pada fungsi pencarian, tampilan relasi, stabilitas build, dan kesesuaian alur Semantic Web.",
        ],
    )

    add_heading(doc, "6. Tinjauan Singkat Teknologi", 1)
    add_matrix(
        doc,
        ["Komponen", "Peran dalam Sistem"],
        [
            ("RDF", "Merepresentasikan entitas anime dan relasinya dalam bentuk subject-predicate-object."),
            ("SPARQL", "Digunakan untuk query demo pada graph lokal dan pencarian eksternal melalui Wikidata."),
            ("React", "Membangun antarmuka interaktif untuk pencarian, detail resource, graph, dan panel analisis."),
            ("Vite", "Menjalankan development server dan proses build frontend secara ringan."),
            ("Tailwind CSS", "Mengatur sistem visual, layout, komponen, dan responsivitas antarmuka."),
            ("Vercel Serverless Function", "Menjalankan API /api/sparql di sisi server agar request SPARQL tidak langsung dilakukan dari browser."),
        ],
        [1900, 7460],
    )

    add_heading(doc, "7. Metodologi Pengembangan", 1)
    add_para(
        doc,
        "Metodologi pengembangan dilakukan secara bertahap, dimulai dari pengumpulan data, pemodelan relasi, sinkronisasi data, pembangunan antarmuka, integrasi API, hingga pengujian. Pendekatan ini dipilih karena proyek Semantic Web membutuhkan konsistensi antara model data, query, dan representasi visual.",
    )
    add_numbered(
        doc,
        [
            "Mengidentifikasi entitas utama yang relevan, yaitu anime, genre, studio, tema, karakter, dan karya terkait.",
            "Membaca data RDF lokal dari file Turtle dan mengekstrak literal serta relasi antar resource.",
            "Memperkaya graph dengan dataset eksternal dari arsip anime lokal, kemudian melakukan deduplikasi judul.",
            "Menghasilkan file animegraph.json yang berisi entitas, relasi, incoming edge, degree, statistik, dan node terhubung teratas.",
            "Membangun fitur frontend untuk semantic search, detail resource, graph neighborhood, grounded QA, semantic compare, dan SPARQL lab.",
            "Membuat API route /api/sparql untuk menjalankan pencarian Wikidata SPARQL secara server-side.",
            "Melakukan build dan pengujian tampilan desktop maupun mobile untuk memastikan aplikasi dapat digunakan dengan baik.",
        ],
    )

    add_heading(doc, "8. Perancangan Sistem", 1)
    add_heading(doc, "8.1 Arsitektur Umum", 2)
    add_para(
        doc,
        "Arsitektur AnimeGraph Nexus terdiri dari tiga lapisan utama: lapisan data, lapisan pemrosesan, dan lapisan presentasi. Lapisan data berisi RDF lokal serta dataset eksternal. Lapisan pemrosesan dijalankan melalui script sinkronisasi yang mengubah RDF dan CSV menjadi JSON graph. Lapisan presentasi dibangun dengan React untuk menampilkan pencarian, detail resource, visualisasi graph, dan hasil query.",
    )
    add_matrix(
        doc,
        ["Lapisan", "Isi", "Output"],
        [
            ("Data", "data.ttl, arsip CSV anime, dan sumber Wikidata", "Triple RDF dan data mentah anime"),
            ("Pemrosesan", "sync-rdf-data.mjs dan normalisasi resource", "animegraph.json"),
            ("Aplikasi", "React, Tailwind, animegraph.js", "UI eksplorasi graph"),
            ("API", "api/sparql.js", "Hasil pencarian Wikidata dalam JSON"),
        ],
        [1700, 4300, 3860],
    )

    add_heading(doc, "8.2 Model Data", 2)
    add_para(
        doc,
        "Model data menggunakan resource dengan prefix ag: sebagai identitas lokal. Setiap entitas memiliki informasi dasar seperti label, entityType, format, genre, releaseYear, description, ragContext, dan source. Relasi antar entitas direpresentasikan melalui predicate yang dapat digunakan untuk membangun graph neighborhood dan perhitungan korelasi.",
    )
    add_matrix(
        doc,
        ["Relasi", "Makna"],
        [
            ("ag:producedBy", "Menghubungkan anime dengan studio atau pihak produksi."),
            ("ag:hasGenre", "Menghubungkan anime dengan genre."),
            ("ag:hasTheme", "Menghubungkan anime dengan tema naratif atau konsep."),
            ("ag:featuresCharacter", "Menghubungkan anime dengan karakter yang relevan."),
            ("ag:relatedTo", "Menghubungkan anime dengan karya atau resource yang berkaitan."),
        ],
        [2300, 7060],
    )

    add_heading(doc, "9. Fitur Utama Sistem", 1)
    add_bullets(
        doc,
        [
            "Semantic Search: pengguna dapat mencari resource berdasarkan judul, tipe, format, genre, deskripsi, dan konteks RDF.",
            "Resource Detail: aplikasi menampilkan fact matrix, genre, studio, tema, karakter, degree, dan sumber data.",
            "Graph Neighborhood: aplikasi menampilkan node pusat dan relasi langsung ke node lain seperti genre, studio, tema, dan related anime.",
            "Semantic Compare: sistem membandingkan dua anime berdasarkan relasi bersama dan menghitung skor kemiripan.",
            "Connected Recommendations: sistem menampilkan kandidat anime berkorelasi berdasarkan relasi semantik yang sama.",
            "Grounded QA: pengguna dapat bertanya tentang studio, tema, karakter, atau rekomendasi, lalu jawaban dibangun dari fakta graph.",
            "SPARQL Lab: aplikasi menyediakan contoh query SPARQL terkurasi untuk memahami cara data graph dapat ditelusuri.",
            "Live Wikidata Search: API server-side menjalankan pencarian anime ke Wikidata SPARQL endpoint dan mengembalikan hasil dalam JSON.",
        ],
    )

    add_heading(doc, "10. Implementasi", 1)
    add_para(
        doc,
        "Implementasi frontend dilakukan menggunakan React, Vite, dan Tailwind CSS. Struktur utama aplikasi berada pada App.jsx, sedangkan logika pencarian, perbandingan, korelasi, dan query demo berada pada animegraph.js. Data yang digunakan oleh UI disediakan dalam public/data/animegraph.json, sehingga aplikasi dapat membaca graph secara langsung saat dijalankan di browser.",
    )
    add_para(
        doc,
        "Pada sisi data, script sync-rdf-data.mjs membaca file Turtle, mengekstrak literal dan relasi, menghitung incoming relation serta degree, lalu membuat ringkasan statistik seperti total entity, relation count, type distribution, format distribution, genre distribution, dan top connected nodes. Berdasarkan build terakhir, graph berisi 14.078 entitas dan 64.649 relasi.",
    )
    add_para(
        doc,
        "Untuk mengatasi risiko pemanggilan SPARQL langsung dari browser, dibuat API route api/sparql.js. Route ini menerima parameter q, membangun query SPARQL yang aman melalui escaping string, mengirim request POST ke Wikidata, menerapkan timeout, mengelompokkan hasil berdasarkan URI, dan mengembalikan data anime dalam bentuk JSON. Strategi ini membuat integrasi SPARQL lebih sesuai untuk deployment serverless.",
    )

    add_heading(doc, "11. Rencana Pengujian", 1)
    add_matrix(
        doc,
        ["Aspek", "Skenario Uji", "Indikator Keberhasilan"],
        [
            ("Build", "Menjalankan npm run build", "Aplikasi berhasil dibuild tanpa error."),
            ("Data", "Menjalankan npm run sync:data", "animegraph.json terbentuk dan memuat entitas serta relasi."),
            ("Search", "Mencari keyword seperti Steins;Gate, Naruto, dan Death Note", "Hasil relevan muncul dan detail resource mengikuti pilihan."),
            ("Graph", "Membuka resource anime", "Neighborhood graph menampilkan node pusat dan relasi langsung."),
            ("Compare", "Membandingkan dua judul anime", "Sistem menampilkan skor dan alasan kemiripan."),
            ("API SPARQL", "Memanggil /api/sparql?q=naruto", "API mengembalikan JSON dengan status ok dan daftar data anime."),
            ("Responsive", "Menguji layout desktop dan mobile", "Tidak ada horizontal overflow dan konten tetap terbaca."),
        ],
        [1600, 3900, 4860],
    )

    add_heading(doc, "12. Jadwal Pengerjaan", 1)
    add_matrix(
        doc,
        ["Tahap", "Kegiatan", "Estimasi"],
        [
            ("1", "Analisis kebutuhan, studi data, dan penentuan model relasi", "Minggu 1"),
            ("2", "Pemodelan RDF, parsing Turtle, dan integrasi dataset eksternal", "Minggu 2"),
            ("3", "Pembuatan JSON graph, indexing, dan logika pencarian", "Minggu 3"),
            ("4", "Pengembangan UI utama, detail resource, graph, dan dashboard", "Minggu 4"),
            ("5", "Pengembangan semantic compare, grounded QA, dan SPARQL lab", "Minggu 5"),
            ("6", "Integrasi API SPARQL server-side dan pengujian deployment", "Minggu 6"),
            ("7", "Penyempurnaan UI, pengujian akhir, dokumentasi, dan presentasi", "Minggu 7"),
        ],
        [1200, 6460, 1700],
    )

    add_heading(doc, "13. Hasil yang Diharapkan", 1)
    add_bullets(
        doc,
        [
            "Tersedianya aplikasi AnimeGraph Nexus yang dapat digunakan untuk mengeksplorasi data anime berbasis graph.",
            "Tersedianya struktur data RDF-derived graph yang memuat ribuan entitas dan relasi semantik.",
            "Pengguna dapat mencari anime, memahami detail resource, melihat relasi, dan membandingkan kemiripan antar anime.",
            "Aplikasi mampu menunjukkan contoh penerapan RDF dan SPARQL dalam kasus nyata yang mudah dipahami.",
            "Proposal, aplikasi, dan dokumentasi dapat digunakan sebagai bahan evaluasi proyek akhir Semantic Web.",
        ],
    )

    add_heading(doc, "14. Kesimpulan", 1)
    add_para(
        doc,
        "AnimeGraph Nexus dirancang sebagai aplikasi eksplorasi anime yang menekankan relasi semantik, bukan sekadar daftar data. Dengan memanfaatkan RDF, SPARQL, dan visualisasi graph, aplikasi ini dapat membantu pengguna memahami hubungan antar anime melalui genre, studio, tema, karakter, dan karya terkait. Implementasi frontend berbasis React serta integrasi API SPARQL server-side menjadikan sistem ini relevan sebagai proyek akhir Semantic Web yang memiliki aspek konseptual, teknis, dan visual.",
    )

    add_heading(doc, "Daftar Pustaka", 1)
    references = [
        "W3C. RDF 1.1 Concepts and Abstract Syntax.",
        "W3C. SPARQL 1.1 Query Language.",
        "Wikidata. Wikidata Query Service and SPARQL endpoint documentation.",
        "React Documentation. Building user interfaces with components.",
        "Vite Documentation. Frontend tooling and build workflow.",
        "Tailwind CSS Documentation. Utility-first CSS framework.",
        "Vercel Documentation. Serverless Functions and deployment workflow.",
    ]
    for ref in references:
        add_para(doc, ref, after=4)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
